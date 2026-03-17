"""
考卷題目 DOCX 導出 — LaTeX → OMML 原生公式
=============================================
使用 latex2mathml 將 LaTeX 轉為 MathML，
再透過 XSLT 轉為 Word OMML，嵌入 python-docx 段落。
"""

import io
import logging
import re
from typing import List, Optional
from lxml import etree

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

import latex2mathml.converter

logger = logging.getLogger(__name__)

# ================================================================
# MathML → OMML XSLT（Word 原生公式格式）
# ================================================================

# Microsoft 提供的 MML2OMML.xsl 精簡版 — 只需核心轉換邏輯
# 完整 XSLT 太大，這裡用 latex2mathml 直接輸出 MathML，
# 再用 python-docx 的 oMath 元素手動構建。
# 替代方案：直接用 Word 的 Unicode math 線性格式。

_LATEX_INLINE = re.compile(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)')
_LATEX_DISPLAY = re.compile(r'\$\$([\s\S]+?)\$\$')


def _mathml_to_omml(mathml_str: str) -> Optional[etree._Element]:
    """將 MathML 字串轉為 Word OMML XML 元素。"""
    try:
        # 移除 xmlns 以便解析
        mathml_str = mathml_str.replace(' xmlns="http://www.w3.org/1998/Math/MathML"', '')
        math_elem = etree.fromstring(mathml_str)

        # 構建 OMML oMathPara 包裝
        # Word 使用 m:oMath 命名空間
        nsmap = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        omath = etree.Element(qn('m:oMath'), nsmap=nsmap)

        # 遞歸轉換 MathML → OMML
        _convert_mathml_node(math_elem, omath, nsmap)
        return omath
    except Exception as e:
        logger.debug("MathML→OMML 轉換失敗: %s", e)
        return None


def _convert_mathml_node(src, parent, nsmap):
    """遞歸將 MathML 節點轉為 OMML 節點。"""
    tag = _local_tag(src)

    if tag == 'math' or tag == 'mrow':
        for child in src:
            _convert_mathml_node(child, parent, nsmap)

    elif tag in ('mn', 'mi', 'mo', 'mtext'):
        # 數字/變量/運算符/文字 → m:r (run)
        r = etree.SubElement(parent, qn('m:r'))
        rpr = etree.SubElement(r, qn('m:rPr'))
        if tag == 'mi':
            # 變量用斜體
            sty = etree.SubElement(rpr, qn('m:sty'))
            sty.set(qn('m:val'), 'i')
        elif tag == 'mtext':
            # 普通文字（如 cm、kg）→ 正常字體，非數學斜體
            nor = etree.SubElement(rpr, qn('m:nor'))
            nor.set(qn('m:val'), '1')
        # 為所有 math run 設定 Word 字體，避免 Cambria Math 缺字
        wrpr = etree.SubElement(r, qn('w:rPr'))
        rfonts = etree.SubElement(wrpr, qn('w:rFonts'))
        rfonts.set(qn('w:ascii'), 'Times New Roman')
        rfonts.set(qn('w:hAnsi'), 'Times New Roman')
        rfonts.set(qn('w:eastAsia'), '微軟正黑體')
        t = etree.SubElement(r, qn('m:t'))
        t.text = src.text or ''
        t.set(qn('xml:space'), 'preserve')

    elif tag == 'msup':
        # 上標 → m:sSup
        children = list(src)
        if len(children) >= 2:
            ssup = etree.SubElement(parent, qn('m:sSup'))
            e_elem = etree.SubElement(ssup, qn('m:e'))
            _convert_mathml_node(children[0], e_elem, nsmap)
            sup = etree.SubElement(ssup, qn('m:sup'))
            _convert_mathml_node(children[1], sup, nsmap)

    elif tag == 'msub':
        # 下標 → m:sSub
        children = list(src)
        if len(children) >= 2:
            ssub = etree.SubElement(parent, qn('m:sSub'))
            e_elem = etree.SubElement(ssub, qn('m:e'))
            _convert_mathml_node(children[0], e_elem, nsmap)
            sub = etree.SubElement(ssub, qn('m:sub'))
            _convert_mathml_node(children[1], sub, nsmap)

    elif tag == 'mfrac':
        # 分數 → m:f
        children = list(src)
        if len(children) >= 2:
            frac = etree.SubElement(parent, qn('m:f'))
            num = etree.SubElement(frac, qn('m:num'))
            _convert_mathml_node(children[0], num, nsmap)
            den = etree.SubElement(frac, qn('m:den'))
            _convert_mathml_node(children[1], den, nsmap)

    elif tag == 'msqrt':
        # 平方根 → m:rad
        rad = etree.SubElement(parent, qn('m:rad'))
        radpr = etree.SubElement(rad, qn('m:radPr'))
        deg_hide = etree.SubElement(radpr, qn('m:degHide'))
        deg_hide.set(qn('m:val'), '1')
        etree.SubElement(rad, qn('m:deg'))
        e_elem = etree.SubElement(rad, qn('m:e'))
        for child in src:
            _convert_mathml_node(child, e_elem, nsmap)

    elif tag == 'mover':
        # 上方標記（如向量箭頭）
        children = list(src)
        if len(children) >= 2:
            # 簡化：直接輸出 base + 上標
            _convert_mathml_node(children[0], parent, nsmap)

    elif tag == 'munder':
        children = list(src)
        if len(children) >= 2:
            _convert_mathml_node(children[0], parent, nsmap)

    elif tag == 'msubsup':
        # 同時上下標
        children = list(src)
        if len(children) >= 3:
            ssubsup = etree.SubElement(parent, qn('m:sSubSup'))
            e_elem = etree.SubElement(ssubsup, qn('m:e'))
            _convert_mathml_node(children[0], e_elem, nsmap)
            sub = etree.SubElement(ssubsup, qn('m:sub'))
            _convert_mathml_node(children[1], sub, nsmap)
            sup = etree.SubElement(ssubsup, qn('m:sup'))
            _convert_mathml_node(children[2], sup, nsmap)

    else:
        # 未知標籤：嘗試遞歸子節點
        for child in src:
            _convert_mathml_node(child, parent, nsmap)
        # 如果有文字內容
        if src.text:
            r = etree.SubElement(parent, qn('m:r'))
            t = etree.SubElement(r, qn('m:t'))
            t.text = src.text
            t.set(qn('xml:space'), 'preserve')


def _local_tag(elem):
    """去除命名空間，只取標籤名。"""
    tag = elem.tag
    if '}' in tag:
        return tag.split('}', 1)[1]
    return tag


# ================================================================
# 段落構建 — 混合文字 + LaTeX 公式
# ================================================================

def _add_mixed_paragraph(doc, text: str, font_size: int = 12, bold: bool = False):
    """
    將包含 LaTeX 的文字添加到文檔。
    $...$ 和 $$...$$ 會被轉為 OMML 公式，其餘為純文字。
    """
    para = doc.add_paragraph()

    # 先處理 display math ($$...$$)，再處理 inline ($...$)
    # 統一替換為佔位符後逐段處理
    segments = _split_text_and_math(text)

    for seg_type, seg_content in segments:
        if seg_type == 'text':
            run = para.add_run(seg_content)
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
            run.font.bold = bold
            # 中文字體
            run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')
        elif seg_type == 'math':
            _insert_omml_in_paragraph(para, seg_content)

    return para


def _split_text_and_math(text: str) -> List[tuple]:
    """
    將文字拆分為 ('text', ...) 和 ('math', ...) 段落。
    支持 $$...$$ (display) 和 $...$ (inline)。
    """
    segments = []
    pos = 0

    # 合併匹配 display 和 inline
    pattern = re.compile(r'\$\$([\s\S]+?)\$\$|\$([^$]+?)\$')

    for m in pattern.finditer(text):
        # 前面的純文字
        if m.start() > pos:
            segments.append(('text', text[pos:m.start()]))

        latex = m.group(1) or m.group(2)
        segments.append(('math', latex.strip()))
        pos = m.end()

    # 剩餘文字
    if pos < len(text):
        segments.append(('text', text[pos:]))

    return segments


def _insert_omml_in_paragraph(para, latex: str):
    """將 LaTeX 轉為 OMML 並插入段落。"""
    try:
        mathml = latex2mathml.converter.convert(latex)
        omml = _mathml_to_omml(mathml)
        if omml is not None:
            para._element.append(omml)
            return
    except Exception as e:
        logger.debug("LaTeX→OMML 失敗 (%s): %s", latex[:30], e)

    # 回退：純文字顯示 LaTeX
    run = para.add_run(f' [{latex}] ')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)


# ================================================================
# 公開 API — 單題導出
# ================================================================

def export_question_to_docx(
    question: dict,
    index: int = 1,
    subject: str = "math",
) -> io.BytesIO:
    """
    將單道題目導出為 .docx（含 OMML 公式）。

    Returns: BytesIO 可直接作為 StreamingResponse 返回。
    """
    doc = Document()

    # 設定默認字體
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # 標題
    heading = doc.add_heading(f'第 {index} 題', level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 題型 + 分數
    q_type = question.get('question_type', '')
    points = question.get('points', '')
    type_map = {
        'multiple_choice': '選擇題',
        'short_answer': '簡答題',
        'long_answer': '解答題',
        'fill_blank': '填空題',
    }
    meta_text = f"【{type_map.get(q_type, q_type)}】"
    if points:
        meta_text += f"  ({points} 分)"
    meta_para = doc.add_paragraph()
    run = meta_para.add_run(meta_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # 題目
    question_text = question.get('question', '')
    _add_mixed_paragraph(doc, question_text, font_size=12)

    # 選項（選擇題）
    options = question.get('options', [])
    if options:
        for j, opt in enumerate(options):
            letter = chr(65 + j)
            _add_mixed_paragraph(doc, f'{letter}. {opt}', font_size=11)

    # 分隔線
    doc.add_paragraph('─' * 40)

    # 答案
    answer = question.get('correct_answer', '')
    if answer:
        ans_heading = doc.add_paragraph()
        run = ans_heading.add_run('參考答案：')
        run.font.bold = True
        run.font.size = Pt(12)
        _add_mixed_paragraph(doc, answer, font_size=12)

    # 評分標準
    marking = question.get('marking_scheme', '')
    if marking:
        mark_heading = doc.add_paragraph()
        run = mark_heading.add_run('評分標準：')
        run.font.bold = True
        run.font.size = Pt(11)
        _add_mixed_paragraph(doc, marking, font_size=11)

    # 輸出
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def export_session_to_docx(
    questions: list,
    subject: str = "math",
    exam_context: str = "",
    total_marks: Optional[int] = None,
) -> io.BytesIO:
    """
    將整份試卷導出為 .docx。
    """
    doc = Document()

    # 默認字體
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')

    # 試卷標題
    title = doc.add_heading('數學試卷', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if exam_context:
        ctx = doc.add_paragraph()
        ctx.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ctx.add_run(exam_context)
        run.font.size = Pt(14)

    if total_marks:
        marks_para = doc.add_paragraph()
        marks_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = marks_para.add_run(f'總分：{total_marks} 分')
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph('')  # 空行

    type_map = {
        'multiple_choice': '選擇題',
        'short_answer': '簡答題',
        'long_answer': '解答題',
        'fill_blank': '填空題',
    }

    for i, q in enumerate(questions):
        q_type = q.get('question_type', '')
        points = q.get('points', '')

        # 題號 + 元數據
        header_text = f'{i + 1}.'
        if points:
            header_text += f' ({points} 分)'
        header_para = doc.add_paragraph()
        run = header_para.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(12)

        # 題目
        question_text = q.get('question', '')
        _add_mixed_paragraph(doc, question_text, font_size=12)

        # 選項
        options = q.get('options', [])
        if options:
            for j, opt in enumerate(options):
                letter = chr(65 + j)
                _add_mixed_paragraph(doc, f'  {letter}. {opt}', font_size=11)

        doc.add_paragraph('')  # 題間空行

    # 答案頁
    doc.add_page_break()
    doc.add_heading('參考答案', level=1)

    for i, q in enumerate(questions):
        answer = q.get('correct_answer', '')
        if answer:
            _add_mixed_paragraph(doc, f'{i + 1}. {answer}', font_size=11)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
