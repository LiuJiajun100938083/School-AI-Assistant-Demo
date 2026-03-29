"""
通告 DOCX 構建器
================
基於學校模板（含校徽 header）生成正式 Word 通告文件。

策略：
1. 從 notice_templates/samples/ 找到同類型模板作為基底（保留 header/校徽/頁面設定）
2. 清空正文段落
3. 按 AI 生成的通告內容逐行寫入，自動識別標題行、列表項等格式
"""

import logging
import os
import re
from datetime import datetime
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

logger = logging.getLogger(__name__)

TEMPLATES_DIR = "notice_templates"
SAMPLES_DIR = os.path.join(TEMPLATES_DIR, "samples")
OUTPUT_DIR = "generated_notices"


def build_notice_docx(
    session_id: str,
    title: str,
    content: str,
    ref_no: Optional[str] = None,
    notice_type: str = "general",
) -> str:
    """
    構建通告 DOCX 文件。

    Args:
        session_id: 會話 ID（用於文件名）
        title: 通告標題
        content: 通告正文（純文本，AI 生成）
        ref_no: 通告編號（可選）
        notice_type: 通告類型（用於選擇模板）

    Returns:
        str: 生成的 DOCX 文件路徑
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # 1. 選擇基底模板（帶校徽 header）
    doc = _load_base_template(notice_type)

    # 2. 清空正文（安全方式：逐一清除 run，保留文檔結構完整性）
    body = doc.element.body
    from docx.oxml.ns import qn
    # 僅移除 body 直屬的段落元素，不動 header/footer/section
    for p_elem in list(body.findall(qn('w:p'))):
        body.remove(p_elem)
    # 同時移除 body 裡的 table（模板表格）
    for tbl_elem in list(body.findall(qn('w:tbl'))):
        body.remove(tbl_elem)

    # 3. 解析 AI 生成的通告內容並寫入
    _write_notice_content(doc, content, title)

    # 4. 保存
    date_str = datetime.now().strftime("%Y-%m-%d")
    filename = f"通告_{date_str}_{session_id[:8]}.docx"
    file_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(file_path)

    logger.info("通告 DOCX 已生成: %s", file_path)
    return file_path


def _load_base_template(notice_type: str) -> Document:
    """
    載入基底模板：優先同類型 sample → blank_template → 空文檔。
    使用 sample 模板可以保留校徽 header 和頁面設定。
    """
    # 嘗試同類型 sample
    type_dir = os.path.join(SAMPLES_DIR, notice_type)
    if os.path.isdir(type_dir):
        for fname in sorted(os.listdir(type_dir)):
            if fname.endswith(".docx"):
                try:
                    return Document(os.path.join(type_dir, fname))
                except Exception as e:
                    logger.warning("載入模板失敗 (%s): %s", fname, e)

    # 嘗試 general 目錄
    general_dir = os.path.join(SAMPLES_DIR, "general")
    if os.path.isdir(general_dir):
        for fname in sorted(os.listdir(general_dir)):
            if fname.endswith(".docx"):
                try:
                    return Document(os.path.join(general_dir, fname))
                except Exception as e:
                    logger.warning("載入 general 模板失敗: %s", e)

    # 嘗試 blank_template
    blank = os.path.join(TEMPLATES_DIR, "blank_template.docx")
    if os.path.exists(blank):
        try:
            return Document(blank)
        except Exception:
            pass

    # 最後回退：空文檔
    return Document()


def _write_notice_content(doc: Document, content: str, title: str):
    """
    解析通告純文本並寫入 Word 文檔，自動識別格式。
    """
    lines = content.strip().split('\n')

    for line in lines:
        stripped = line.strip()

        # 跳過空行但保留間距
        if not stripped:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=0)
            continue

        # 跳過 markdown 分隔線
        if stripped in ('---', '===', '***'):
            continue

        # 清理 markdown 格式
        cleaned = _clean_markdown(stripped)

        # 判斷行類型並設定格式
        if _is_title_line(cleaned, title):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cleaned)
            run.bold = True
            run.font.size = Pt(16)
            _set_paragraph_spacing(p, before=Pt(6), after=Pt(6))

        elif cleaned.startswith('敬啟者') or cleaned.startswith('敬启者'):
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            _set_paragraph_spacing(p, before=Pt(6), after=Pt(6))

        elif cleaned.startswith('此致'):
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            _set_paragraph_spacing(p, before=Pt(12), after=Pt(2))

        elif _is_signature_line(cleaned):
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            _set_paragraph_spacing(p, before=Pt(2), after=Pt(2))

        elif _is_field_line(cleaned):
            # 結構化字段行（日期：xxx、時間：xxx）
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            _set_paragraph_spacing(p, before=Pt(2), after=Pt(2))

        else:
            # 普通正文段落
            p = doc.add_paragraph()
            # 首行縮進效果（中文排版習慣）
            if not cleaned.startswith('　') and len(cleaned) > 20:
                cleaned = '　　' + cleaned
            run = p.add_run(cleaned)
            _set_paragraph_spacing(p, before=Pt(2), after=Pt(2))


def _clean_markdown(text: str) -> str:
    """清理 markdown 格式標記"""
    # 去除 **bold** 標記
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 去除 *italic* 標記
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 去除列表標記 (- 或 *)
    text = re.sub(r'^[\-\*]\s+', '', text)
    # 去除數字列表標記
    text = re.sub(r'^\d+[\.\)]\s+', '', text)
    return text.strip()


def _is_title_line(text: str, title: str) -> bool:
    """判斷是否為標題行"""
    if '通知' in text and len(text) < 30:
        return True
    if '通告' in text and len(text) < 30:
        return True
    if text.startswith('【') and text.endswith('】'):
        return True
    if title and title in text and len(text) < 40:
        return True
    return False


def _is_signature_line(text: str) -> bool:
    """判斷是否為簽名相關行"""
    keywords = ['校長', '學生家長', '貴家長', '二零', '家長', '副校長']
    return any(k in text for k in keywords) and len(text) < 40


def _is_field_line(text: str) -> bool:
    """判斷是否為結構化字段行（日期：、時間：等）"""
    field_prefixes = [
        '日', '時間', '地點', '對象', '費用', '考試',
        '活動', '會議', '注意', '聯絡', '如有',
    ]
    return any(text.startswith(p) for p in field_prefixes) and '：' in text[:10]


def _set_paragraph_spacing(p, before=None, after=None):
    """設定段落間距"""
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = before
    if after is not None:
        pf.space_after = after
