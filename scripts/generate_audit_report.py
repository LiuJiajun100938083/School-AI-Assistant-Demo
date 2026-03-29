"""Generate the architecture audit report as a DOCX document."""
import sys, os
sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Helpers ──

def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None, header_color="006633"):
    """Add a styled table with colored header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F5F5F7")

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    return table

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    return h

def add_code_block(doc, code_text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Consolas"/>')
    rPr.append(rFonts)
    return p

def add_severity_tag(paragraph, severity):
    """Add a colored severity tag inline."""
    colors = {
        "CRITICAL": "CC0000",
        "HIGH": "E65100",
        "MEDIUM": "F59E0B",
        "LOW": "6B7280",
    }
    color = colors.get(severity, "333333")
    run = paragraph.add_run(f"  [{severity}]")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)

# ── Build Document ──

doc = Document()

# -- Page setup --
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# -- Default font --
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ================================================================
# TITLE PAGE
# ================================================================

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("School-AI-Assistant-Demo")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("全项目架构审查报告")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("审查日期：2026-02-23\n审查标准：企业级工程规范\n审查范围：全部后端 + 前端代码")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

doc.add_paragraph()
doc.add_paragraph()

grade_p = doc.add_paragraph()
grade_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = grade_p.add_run("综合评级：B-")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

grade_desc = doc.add_paragraph()
grade_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = grade_desc.add_run("良好骨架，局部失控")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

doc.add_page_break()

# ================================================================
# TABLE OF CONTENTS (manual)
# ================================================================

add_heading(doc, "目录", level=1)
toc_items = [
    "1. 项目概况",
    "2. 架构总体评估",
    "3. CRITICAL 严重问题（6 个）",
    "4. HIGH 高优问题（8 个）",
    "5. MEDIUM 中等问题（12 个）",
    "6. LOW 低优问题（8 个）",
    "7. 模块健康度评分",
    "8. 优先修复路线图",
    "9. 总结与建议",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(1)
    for run in p.runs:
        run.font.size = Pt(11)

doc.add_page_break()

# ================================================================
# 1. PROJECT OVERVIEW
# ================================================================

add_heading(doc, "1. 项目概况", level=1)

add_styled_table(doc,
    ["指标", "数值"],
    [
        ["Python 后端代码", "~45,851 行（95+ 文件）"],
        ["前端代码 (JS/CSS/HTML)", "~61,889 行（47+ 文件）"],
        ["业务领域", "15 个（认证、课堂、考勤、错题本、AI学习中心、论坛等）"],
        ["路由文件", "17 个"],
        ["数据库迁移文件", "10+ 个"],
        ["测试文件", "2 个（仅数据库连接测试）"],
    ],
    col_widths=[5, 12],
)

doc.add_paragraph()
add_heading(doc, "后端文件规模 Top 10", level=2)

add_styled_table(doc,
    ["文件", "行数", "职责"],
    [
        ["routers/attendance.py", "4,108", "考勤系统（跨层违规）"],
        ["domains/mistake_book/service.py", "1,664", "错题本业务逻辑"],
        ["domains/ai_learning_center/service.py", "1,167", "AI学习中心业务"],
        ["routers/china_game.py", "1,121", "经济模拟游戏（零认证）"],
        ["routers/game_upload.py", "1,078", "游戏上传（内嵌服务）"],
        ["routers/ai_learning_center.py", "979", "学习中心路由"],
        ["domains/classroom/service.py", "959", "课堂管理业务"],
        ["domains/attendance/service.py", "908", "考勤业务（未被使用）"],
        ["domains/vision/service.py", "725", "图像识别服务"],
        ["domains/chat/service.py", "630", "聊天/AI对话"],
    ],
    col_widths=[6.5, 1.5, 9],
)

doc.add_paragraph()
add_heading(doc, "前端文件规模", level=2)

add_styled_table(doc,
    ["文件", "行数", "说明"],
    [
        ["js/ai_learning_center.js", "~4,015", "学习中心全部 JS 逻辑（单体）"],
        ["css/ai_learning_center.css", "~3,550", "学习中心样式"],
        ["ai_learning_center.html", "~675", "学习中心 HTML"],
        ["js/attendance.js", "~2,500", "考勤前端"],
        ["js/china_economy_game.js", "~2,000", "经济游戏前端"],
    ],
    col_widths=[6.5, 1.5, 9],
)

doc.add_page_break()

# ================================================================
# 2. ARCHITECTURE ASSESSMENT
# ================================================================

add_heading(doc, "2. 架构总体评估", level=1)

doc.add_paragraph("项目主体采用四层分层架构：")
add_code_block(doc,
    "理想分层:  UI (HTML/JS)  →  Router  →  Service  →  Repository  →  Database\n"
    "                                  ↑\n"
    "                               Config / DI Container"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("合规模块（11/14）：")
run.bold = True
doc.add_paragraph(
    "classroom, chat, user, learning_task, subject, teacher_class, analytics, "
    "notice, auth, learning_modes, ai_learning_center — 均遵循 Router → Service → Repository 分层，"
    "通过 ServiceContainer 依赖注入。"
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("违规模块（3/14）：")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

violations = [
    ("attendance.py", "路由层直接执行 SQL，无 Service/Repository 层，4,108 行耦合"),
    ("game_upload.py", "路由文件内嵌 3 个完整类（Config + Service + DBService），认证绕过"),
    ("china_game.py", "独立子系统，全部端点零认证，模块级 Service 实例化"),
]
for name, desc in violations:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"{name}")
    run.bold = True
    p.add_run(f" — {desc}")

doc.add_page_break()

# ================================================================
# 3. CRITICAL ISSUES
# ================================================================

add_heading(doc, "3. CRITICAL 严重问题（6 个）", level=1)

# C1
p = add_heading(doc, "C1. attendance.py — 路由层直接执行 SQL", level=2)
doc.add_paragraph(
    "文件：app/routers/attendance.py（4,108 行，全项目最大单文件）\n"
    "问题：整个考勤系统绕过 Service/Repository 层，在路由 handler 中直接 pool.execute() 执行原始 SQL。"
)
doc.add_paragraph("违规项：")
for item in [
    "跨层违规 — Router 层直接访问数据库",
    "单一职责违反 — HTTP 处理、业务逻辑、数据访问全部耦合",
    "不可测试 — 业务逻辑无法脱离 HTTP 上下文单独测试",
    "事务管理缺失 — 无统一事务控制",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph("影响：考勤是核心业务，此模块完全不可维护、不可扩展。")

doc.add_paragraph()

# C2
add_heading(doc, "C2. game_upload.py — 认证绕过漏洞", level=2)
doc.add_paragraph(
    "文件：app/routers/game_upload.py（1,078 行）\n"
    "问题：路由文件内定义了 GameUploadConfig、GameUploadService、GameDBService 三个完整类。"
    "更严重的是，接受 form data 中的 user_id 和 user_role，任何客户端可伪造身份。"
)
add_code_block(doc,
    "# 客户端可以发送 user_role=\"admin\" 来获取管理员权限\n"
    "user_id: str = Form(...)\n"
    "user_role: str = Form(...)"
)
p = doc.add_paragraph()
run = p.add_run("安全等级：极高风险。攻击者可冒充任何用户和角色。")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_paragraph()

# C3
add_heading(doc, "C3. china_game.py — 全部端点零认证", level=2)
doc.add_paragraph(
    "文件：app/routers/china_game.py（1,121 行）\n"
    "问题：所有游戏 API 端点无需登录即可调用，无任何 Depends(get_current_user)。"
    "任何未认证用户可完全访问游戏功能。"
)

doc.add_paragraph()

# C4
add_heading(doc, "C4. security.py — 内存级安全状态", level=2)
doc.add_paragraph(
    "文件：app/core/security.py\n"
    "问题：\n"
    "- Token 黑名单 = Python set()，服务重启即丢失，已注销的 token 重新生效\n"
    "- 登录尝试计数 = Python dict()，多 worker 进程不共享\n"
    "- 生产环境下暴力破解限速功能完全失效"
)
doc.add_paragraph("建议：改用 Redis 或数据库表存储黑名单和限速状态。")

doc.add_paragraph()

# C5
add_heading(doc, "C5. 前端 XSS 漏洞 — marked.parse() 无清洗", level=2)
doc.add_paragraph(
    "文件：web_static/js/ai_learning_center.js（行 846, 2510-2513）\n"
    "问题：Markdown 内容通过 marked.parse() 生成 HTML 后直接赋值 innerHTML，无 DOMPurify 清洗。"
    "如果管理员输入的 Markdown 包含 <script> 标签或事件处理器，会直接执行。"
)
add_code_block(doc,
    "// 行 846 - article 内容直接渲染\n"
    "bodyEl.innerHTML = `<div>${marked.parse(articleContent)}</div>`;\n\n"
    "// 修复方案：\n"
    "bodyEl.innerHTML = `<div>${DOMPurify.sanitize(marked.parse(articleContent))}</div>`;"
)

doc.add_paragraph()

# C6
add_heading(doc, "C6. 测试覆盖率约等于 0%", level=2)
doc.add_paragraph(
    "全项目仅 2 个测试文件（database_migration/ 下的连接测试），核心业务逻辑零单元测试。\n"
    "45,000+ 行后端代码、4,000 行前端 JS 均无自动化测试保护。\n"
    "任何重构或修改都可能引入回归 bug 而无法被发现。"
)

doc.add_page_break()

# ================================================================
# 4. HIGH ISSUES
# ================================================================

add_heading(doc, "4. HIGH 高优问题（8 个）", level=1)

high_issues = [
    (
        "H1. 认证辅助函数跨 7+ 路由重复",
        "attendance.py, china_game.py, game_upload.py 等各自实现 get_current_user_from_cookie()，"
        "而 app/core/dependencies.py 已有标准实现。违反 DRY 原则。"
    ),
    (
        "H2. 超长函数（远超 30 行标准）",
        None  # special: table
    ),
    (
        "H3. 考勤系统 scan/manual_scan 大量重复",
        "app/routers/attendance.py 中 manual_scan_attendance 与 scan_attendance 逻辑 80%+ 重复，"
        "仅输入来源不同。应提取为共享 Service 方法。"
    ),
    (
        "H4. 前端管理后台仍有 inline onclick",
        "ai_learning_center.js 中 loadAdminTab() 渲染管理列表时仍使用 inline onclick。"
        "节点详情面板已修复，但 admin 列表的编辑/删除按钮尚未迁移到事件委托。"
    ),
    (
        "H5. mistake_book 路由跨层实例化 Repository",
        "app/routers/mistake_book.py 直接 MistakeBookRepository() 而非通过 ServiceContainer 获取。"
        "违反依赖注入原则。"
    ),
    (
        "H6. renderKnowledgeMap() 嵌套深度达 6 层",
        "ai_learning_center.js 行 ~1100-1600，tick 回调内嵌套：forEach → if → if → ternary → callback。"
        "远超 3 层标准，严重影响可读性。"
    ),
    (
        "H7. AI 模型名硬编码",
        "llm/providers/ollama.py 和 domains/chat/service.py 中 \"gemma3:12b\" 等模型名直接写在代码中，"
        "应提取到配置文件或环境变量。"
    ),
    (
        "H8. BaseRepository f-string 拼接表名",
        "app/infrastructure/database/base_repository.py 中表名通过 f-string 插入 SQL。"
        "虽然表名来自代码内部，但违反参数化查询原则，且不利于 SQL 审计。"
    ),
]

for title, desc in high_issues:
    add_heading(doc, title, level=2)
    if title.startswith("H2"):
        # Special: function length table
        doc.add_paragraph("以下函数严重超出 30 行标准：")
        add_styled_table(doc,
            ["文件", "函数", "行数", "标准"],
            [
                ["ai_learning_center.js", "renderKnowledgeMap()", "506", "30"],
                ["ai_learning_center.js", "init()", "180", "30"],
                ["ai_learning_center.js", "loadAdminTab()", "170", "30"],
                ["ai_learning_center.js", "showNodeDetail()", "115", "30"],
                ["ai_learning_center.js", "showEbookContent()", "103", "30"],
                ["attendance.py", "多个 scan 函数", "150+", "30"],
                ["mistake_book/service.py", "generate_analysis()", "120+", "30"],
                ["classroom/service.py", "多个方法", "80+", "30"],
            ],
            col_widths=[5, 4.5, 1.5, 1.5],
            header_color="E65100",
        )
    else:
        doc.add_paragraph(desc)
    doc.add_paragraph()

doc.add_page_break()

# ================================================================
# 5. MEDIUM ISSUES
# ================================================================

add_heading(doc, "5. MEDIUM 中等问题（12 个）", level=1)

medium_issues = [
    ("M1", "魔法值泛滥",
     "security.py 中 MAX_LOGIN_ATTEMPTS = 5 未提取到 config；JS 中 API_BASE、debounce 延迟、"
     "opacity 值等散落各处；attendance.py 状态码字符串 \"present\"/\"late\"/\"absent\" 未定义为常量。"),
    ("M2", "Pydantic v1 .dict() 调用（已废弃）",
     "classroom/schemas.py、analytics/models.py、mistake_book/schemas.py 使用已废弃的 .dict()，"
     "应改为 .model_dump()。"),
    ("M3", "N+1 查询模式",
     "ai_learning_center/service.py 行 428-432：遍历 87 个节点时逐个查询 contents，"
     "产生 87 次额外 SQL。应改为单次 JOIN 查询。"),
    ("M4", "前端全局状态无保护",
     "ai_learning_center.js 中 state 对象可被任意函数直接修改，无 getter/setter 封装。"),
    ("M5", "模块级 Service 实例化绕过 DI",
     "china_game.py 和 game_upload.py 在模块顶层直接 service = GameService()，"
     "绕过了 ServiceContainer 依赖注入。"),
    ("M6", "错误信息泄露",
     "多个路由中 except Exception as e: return error_response(\"SERVER_ERROR\", str(e))，"
     "将内部堆栈和数据库错误暴露给客户端。"),
    ("M7", "CSS !important 滥用",
     "ai_learning_center.css 有 14 处 !important，多数可通过提高选择器特异性替代。"),
    ("M8", "前端 fetch 无超时机制",
     "ai_learning_center.js 的 api() 函数所有 fetch() 调用无 AbortController 超时。"),
    ("M9", "D3 事件监听器未清理",
     "renderKnowledgeMap() 重复调用时 SVG 及其事件监听器未销毁旧实例，可能导致内存泄漏。"),
    ("M10", "SQL 迁移文件散落",
     "根目录 4 个 .sql + database_migration/ 4 个 + forum_system/ 2 个，无统一迁移管理工具。"),
    ("M11", "CSS 死代码",
     "旧类名 .alc-panel-content、.alc-linked-contents 等在 HTML 初始内容中仍引用但 CSS 中已无定义。"),
    ("M12", "前端 API 调用无统一错误处理",
     "每个页面各自实现 api() 函数，虽有 shared/api.js 但 ai_learning_center.js 内部又重新实现了一个。"),
]

for code, title, desc in medium_issues:
    p = doc.add_paragraph()
    run = p.add_run(f"{code}. {title}")
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph(desc)
    doc.add_paragraph()

doc.add_page_break()

# ================================================================
# 6. LOW ISSUES
# ================================================================

add_heading(doc, "6. LOW 低优问题（8 个）", level=1)

add_styled_table(doc,
    ["#", "问题", "位置"],
    [
        ["L1", "单字母变量 n, e, d, c, m", "JS D3 回调中大量使用"],
        ["L2", "print() 代替 logger", "scripts/*.py"],
        ["L3", "函数内部 import", "repository.py (import json as _json)"],
        ["L4", "响应格式不统一", "部分路由返回 dict，部分用 success_response()"],
        ["L5", "HTML 语义化不足", "<div> 代替 <nav>, <section> 等"],
        ["L6", "无 ARIA 无障碍属性", "图谱 SVG 节点、弹出面板"],
        ["L7", "CSS 命名混合", "旧 .alc-node-section vs 新 BEM .alc-nd__section"],
        ["L8", "无 .env.example 文件", "新开发者不知需要哪些环境变量"],
    ],
    col_widths=[1, 6, 10],
    header_color="6B7280",
)

doc.add_page_break()

# ================================================================
# 7. MODULE HEALTH SCORES
# ================================================================

add_heading(doc, "7. 模块健康度评分", level=1)

doc.add_paragraph("评分标准：A=优秀 B=良好 C=合格 D=需改进 F=严重不合格")
doc.add_paragraph()

health_data = [
    ["classroom",          "A", "B", "A", "C", "B+"],
    ["chat",               "A", "B", "A", "C", "B+"],
    ["user",               "A", "B", "A", "C", "B"],
    ["learning_task",      "A", "B", "A", "C", "B"],
    ["ai_learning_center", "A", "C", "B", "D", "B-"],
    ["mistake_book",       "B", "C", "B", "D", "C+"],
    ["analytics",          "A", "B", "B", "D", "B-"],
    ["forum_system",       "A", "B", "B", "D", "B"],
    ["attendance",         "F", "D", "C", "F", "D"],
    ["game_upload",        "F", "D", "F", "F", "F"],
    ["china_game",         "C", "C", "F", "F", "D-"],
    ["前端 JS (整体)",       "-", "D", "C", "F", "D+"],
]

table = add_styled_table(doc,
    ["模块", "分层", "代码质量", "安全性", "可测试", "综合"],
    health_data,
    col_widths=[4, 1.8, 2.2, 2, 2, 1.8],
)

# Color-code the grade cells
grade_colors = {
    "A": "E8F5EC", "B+": "E8F5EC", "B": "E8F5EC", "B-": "FFF8E1",
    "C+": "FFF8E1", "C": "FFF8E1", "D+": "FFF3E0", "D": "FFF3E0",
    "D-": "FFEBEE", "F": "FFEBEE", "-": "F5F5F5",
}
for r_idx, row_data in enumerate(health_data):
    row = table.rows[r_idx + 1]
    for c_idx in range(1, 6):
        val = row_data[c_idx]
        color = grade_colors.get(val, "FFFFFF")
        set_cell_shading(row.cells[c_idx], color)

doc.add_page_break()

# ================================================================
# 8. REMEDIATION ROADMAP
# ================================================================

add_heading(doc, "8. 优先修复路线图", level=1)

# P0
add_heading(doc, "P0 — 安全漏洞（立即修复）", level=2)
p = doc.add_paragraph()
run = p.add_run("预计总工时：4.5 小时")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_styled_table(doc,
    ["#", "任务", "工作量"],
    [
        ["1", "game_upload.py 认证绕过：删除 form data 的 user_id/user_role，改用 Depends(get_current_user)", "1h"],
        ["2", "china_game.py 所有端点加 Depends(get_current_user)", "1h"],
        ["3", "前端 marked.parse() 输出加 DOMPurify.sanitize()", "0.5h"],
        ["4", "security.py Token 黑名单改 Redis 或数据库表", "2h"],
    ],
    col_widths=[0.8, 13, 1.8],
    header_color="CC0000",
)

doc.add_paragraph()

# P1
add_heading(doc, "P1 — 架构债务（1-2 周）", level=2)
p = doc.add_paragraph()
run = p.add_run("预计总工时：8-10 个工作日")
run.bold = True
run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

add_styled_table(doc,
    ["#", "任务", "工作量"],
    [
        ["5", "attendance.py 拆分为 Service + Repository 三层架构", "3-5 天"],
        ["6", "game_upload.py 内嵌类提取到 domains/game_upload/", "1 天"],
        ["7", "renderKnowledgeMap() 拆分为 8 个子函数（每个 ≤30 行）", "1 天"],
        ["8", "ai_learning_center.js 拆分为 kg.js / admin.js / viewer.js / ai-chat.js", "2 天"],
        ["9", "魔法值提取为 CONFIG / CONSTANTS 对象", "0.5 天"],
    ],
    col_widths=[0.8, 13, 1.8],
    header_color="E65100",
)

doc.add_paragraph()

# P2
add_heading(doc, "P2 — 质量提升（持续改进）", level=2)
p = doc.add_paragraph()
run = p.add_run("预计总工时：5-7 个工作日")
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)

add_styled_table(doc,
    ["#", "任务", "工作量"],
    [
        ["10", "find_by_node N+1 优化为单次 JOIN 查询", "2h"],
        ["11", "前端 api() 加 AbortController 超时", "1h"],
        ["12", "统一错误响应格式，隐藏内部错误细节", "2h"],
        ["13", ".env.example + 数据库迁移统一到 database_migration/", "1h"],
        ["14", "核心 Service 方法添加单元测试（至少 auth, classroom）", "3-5 天"],
        ["15", "移除 Pydantic .dict() 改 .model_dump()", "0.5h"],
    ],
    col_widths=[0.8, 13, 1.8],
    header_color="006633",
)

doc.add_page_break()

# ================================================================
# 9. SUMMARY
# ================================================================

add_heading(doc, "9. 总结与建议", level=1)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("综合评级：B-（良好骨架，局部失控）")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

add_heading(doc, "做得好的方面", level=2)
for item in [
    "主体架构分层清晰：11/14 个模块遵循 Router → Service → Repository 四层架构",
    "依赖注入：ServiceContainer 统一管理服务实例，避免全局变量",
    "设计系统完善：CSS 变量体系（颜色、间距、圆角、阴影）统一且一致",
    "错误处理框架：自定义异常体系（AppException、ValidationError 等）层次分明",
    "安全基础设施：JWT 认证、密码哈希、CORS 配置、权限检查均已实现",
    "知识图谱可视化：D3.js 力导向图、层级布局、LOD 缩放、展开折叠功能完整",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_heading(doc, "最需要改善的方面", level=2)
for item in [
    "3 个子系统完全脱离架构约束（attendance、game_upload、china_game）",
    "前端 JS 单体文件 4000 行，缺乏模块化拆分",
    "测试覆盖率为零，任何重构都存在回归风险",
    "认证绕过和零认证端点是需要立即修复的安全漏洞",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
add_heading(doc, "建议执行顺序", level=2)
doc.add_paragraph(
    "1. 本周：修复 P0 安全漏洞（4.5 小时即可完成）\n"
    "2. 下两周：拆分 attendance.py 和 game_upload.py（还清最大架构债务）\n"
    "3. 持续：每次新功能开发时同步补充单元测试\n"
    "4. 持续：前端 JS 按功能域逐步拆分为独立模块"
)

doc.add_paragraph()
doc.add_paragraph()

# Footer note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— 报告结束 —")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xA1, 0xA1, 0xA6)

# ================================================================
# SAVE
# ================================================================

output_path = r"C:\Users\15821\School-AI-Assistant-Demo\docs\架构审查报告_2026-02-23.docx"
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Report saved to: {output_path}")
